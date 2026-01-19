"""Word document export utility for NOVA reports."""
import re
from pathlib import Path
from typing import Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def export_to_word(
    content: str,
    company_name: str,
    question: str,
    output_dir: Optional[Path] = None,
    filename: Optional[str] = None
) -> Path:
    """
    Export NOVA report to Word document.
    
    Args:
        content: Markdown-formatted report content
        company_name: Company name for title
        question: Original question
        output_dir: Output directory (default: ./reports)
        filename: Output filename (default: auto-generated)
    
    Returns:
        Path to created Word document
    """
    if not HAS_DOCX:
        raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    # Create output directory
    if output_dir is None:
        output_dir = Path.cwd() / "reports"
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate filename
    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_company = re.sub(r'[^\w\s-]', '', company_name).strip().replace(' ', '_')
        filename = f"NOVA_Report_{safe_company}_{timestamp}.docx"
    
    output_path = output_dir / filename
    
    # Create document
    doc = Document()
    
    # Set document styles
    setup_document_styles(doc)
    
    # Add title
    title = doc.add_heading(f'{company_name} - IR 분석 보고서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    meta_para = doc.add_paragraph()
    meta_para.add_run(f'생성일시: {datetime.now().strftime("%Y년 %m월 %d일 %H:%M")}').font.size = Pt(10)
    meta_para.add_run('\n질문: ').font.size = Pt(10)
    meta_para.add_run(question).font.size = Pt(10)
    meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()  # Spacing
    
    # Parse and add content
    parse_markdown_to_word(doc, content)
    
    # Save document
    doc.save(str(output_path))
    
    return output_path


def setup_document_styles(doc: Document):
    """Setup document styles."""
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)
    
    # Set heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = '맑은 고딕'
        heading_font.bold = True
        if i == 1:
            heading_font.size = Pt(18)
        elif i == 2:
            heading_font.size = Pt(14)
        else:
            heading_font.size = Pt(12)


def parse_markdown_to_word(doc: Document, markdown: str):
    """Parse markdown content and add to Word document."""
    lines = markdown.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Headers
        if line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=1)
        elif line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=2)
        elif line.startswith('#### '):
            text = line[5:].strip()
            doc.add_heading(text, level=3)
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Check for bold
            para = doc.add_paragraph(style='List Bullet')
            add_formatted_text(para, text)
        # Reference links (출처 : URL)
        elif line.startswith('출처 : '):
            url = line[7:].strip()
            para = doc.add_paragraph()
            run = para.add_run('출처 : ')
            run.font.size = Pt(9)
            
            # Add hyperlink
            try:
                add_hyperlink(para, url, url)
            except Exception:
                # Fallback: add as blue underlined text
                run = para.add_run(url)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
                run.font.underline = True
        # Regular paragraph
        else:
            para = doc.add_paragraph()
            add_formatted_text(para, line)
        
        i += 1


def add_formatted_text(paragraph, text: str):
    """Add text with formatting (bold, italic, etc.)."""
    # Simple markdown parsing
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_hyperlink(paragraph, url: str, text: str):
    """Add a hyperlink to a paragraph."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    
    # Create hyperlink relationship
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create hyperlink XML element
    hyperlink_xml = (
        '<w:hyperlink r:id="%s" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:r><w:rPr><w:color w:val="0000FF"/><w:u w:val="single"/></w:rPr>'
        '<w:t xml:space="preserve">%s</w:t></w:r></w:hyperlink>' % (r_id, text)
    )
    
    # Parse and add hyperlink
    hyperlink = parse_xml(hyperlink_xml)
    paragraph._p.append(hyperlink)
